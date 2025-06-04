import sys
import traceback
from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout,
                             QPushButton, QTextEdit, QFileDialog, QLabel)
# from PyQt5.QtCore import Qt # Убрали неиспользуемый импорт
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm

# --- Требования к форматированию ---
TARGET_FONT = "Times New Roman"
TARGET_SIZE_PT = 14.0
TARGET_FIRST_LINE_INDENT_CM = 1.25
TARGET_LINE_SPACING_RULE = WD_LINE_SPACING.ONE_POINT_FIVE

TARGET_MARGIN_TOP_CM = 2.0
TARGET_MARGIN_BOTTOM_CM = 2.0
TARGET_MARGIN_LEFT_CM = 3.0
TARGET_MARGIN_RIGHT_CM = 1.5


class DocFormatChecker(QWidget):
    def __init__(self):
        super().__init__()
        self.setup_ui()

    def _escape_html(self, text_to_escape):
        if text_to_escape is None: return ""
        text_str = str(text_to_escape)
        return text_str.replace('&', '&').replace('<', '<').replace('>', '>').replace('\n', '<br/>').replace('\t',
                                                                                                             '    ')

    def setup_ui(self):
        self.setWindowTitle('Проверка форматирования Word')
        self.setGeometry(300, 300, 750, 700)
        layout = QVBoxLayout()
        self.info_lbl = QLabel('Выберите .docx файл для проверки:', self)
        layout.addWidget(self.info_lbl)
        self.open_btn = QPushButton('Открыть документ...', self)
        self.open_btn.clicked.connect(self.select_file_and_analyze)
        layout.addWidget(self.open_btn)
        self.results_text_edit = QTextEdit(self)
        self.results_text_edit.setReadOnly(True)
        self.results_text_edit.setAcceptRichText(True)
        self.results_text_edit.setLineWrapMode(QTextEdit.WidgetWidth)
        layout.addWidget(self.results_text_edit)
        self.setLayout(layout)
        self.show()

    def select_file_and_analyze(self):
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(self, "Выбрать документ Word", "",
                                                   "Документы Word (*.docx);;Все файлы (*.*)", options=options)
        if file_path:
            self.results_text_edit.clear()
            self.results_text_edit.append(f"<b>Файл:</b> {self._escape_html(file_path)}")
            self.analyze_docx(file_path)
        else:
            self.results_text_edit.setText("Файл не выбран, анализ отменен.")

    def _get_actual_font_for_run(self, run, para_style):
        font_name = run.font.name
        font_size = run.font.size.pt if run.font.size else None
        if para_style:
            if font_name is None and para_style.font and para_style.font.name:
                font_name = para_style.font.name
            if font_size is None and para_style.font and para_style.font.size:
                font_size = para_style.font.size.pt
        return font_name, font_size

    def _check_font_and_size(self, run_text_segment, font_name, font_size_pt):
        errors = []
        if not run_text_segment.strip(): return errors

        if font_name is None:
            errors.append(f"неверный шрифт (должен быть '{TARGET_FONT}')")
        elif font_name != TARGET_FONT:
            errors.append(f"шрифт \"{font_name}\" (должен быть \"{TARGET_FONT}\")")

        if font_size_pt is None:
            errors.append(f"неверный размер (должен быть {TARGET_SIZE_PT:.1f} пт)")
        elif abs(font_size_pt - TARGET_SIZE_PT) > 0.1:
            errors.append(f"размер {font_size_pt:.1f} пт (должен быть {TARGET_SIZE_PT:.1f} пт)")
        return errors

    def _check_paragraph_formatting(self, para):
        para_errors = []
        para_format = para.paragraph_format
        target_indent_emu = Cm(TARGET_FIRST_LINE_INDENT_CM)
        current_indent_emu = para_format.first_line_indent
        if current_indent_emu is None:
            if TARGET_FIRST_LINE_INDENT_CM != 0:
                para_errors.append(f"отступ первой строки не задан (нужен {TARGET_FIRST_LINE_INDENT_CM:.2f} см)")
        elif abs(current_indent_emu.cm - TARGET_FIRST_LINE_INDENT_CM) > 0.05:
            para_errors.append(
                f"отступ первой строки {current_indent_emu.cm:.2f} см (нужен {TARGET_FIRST_LINE_INDENT_CM:.2f} см)")

        current_line_spacing_rule = para_format.line_spacing_rule
        current_line_spacing_value = para_format.line_spacing
        is_target_one_point_five = (TARGET_LINE_SPACING_RULE == WD_LINE_SPACING.ONE_POINT_FIVE)
        current_is_actually_one_point_five = False
        if current_line_spacing_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
            current_is_actually_one_point_five = True
        elif current_line_spacing_rule == WD_LINE_SPACING.MULTIPLE and \
                current_line_spacing_value is not None and \
                abs(current_line_spacing_value - 1.5) < 0.01:
            current_is_actually_one_point_five = True

        if is_target_one_point_five:
            if not current_is_actually_one_point_five:
                rule_name_str = str(current_line_spacing_rule)
                for name, member in WD_LINE_SPACING.__members__.items():
                    if member.value == current_line_spacing_rule: rule_name_str = name; break
                para_errors.append(
                    f"межстрочный интервал не полуторный (тек. правило: {rule_name_str}, знач: {current_line_spacing_value})")
        else:
            if current_line_spacing_rule != TARGET_LINE_SPACING_RULE:
                current_rule_name_str, target_rule_name_str = str(current_line_spacing_rule), str(
                    TARGET_LINE_SPACING_RULE)
                for name, member in WD_LINE_SPACING.__members__.items():
                    if member.value == current_line_spacing_rule: current_rule_name_str = name
                    if member.value == TARGET_LINE_SPACING_RULE: target_rule_name_str = name
                para_errors.append(
                    f"неверное правило межстрочного интервала (тек: \"{current_rule_name_str}\", нужно: \"{target_rule_name_str}\")")
        return para_errors

    def _check_document_margins(self, doc):
        margin_errors = []
        if not doc.sections:
            margin_errors.append("Не удалось проверить поля: в документе нет секций.")
            return margin_errors
        section = doc.sections[0]
        tolerance_cm = 0.05
        if abs(section.top_margin.cm - TARGET_MARGIN_TOP_CM) > tolerance_cm:
            margin_errors.append(f"верхнее поле {section.top_margin.cm:.2f} см (нужно {TARGET_MARGIN_TOP_CM:.2f} см)")
        if abs(section.bottom_margin.cm - TARGET_MARGIN_BOTTOM_CM) > tolerance_cm:
            margin_errors.append(
                f"нижнее поле {section.bottom_margin.cm:.2f} см (нужно {TARGET_MARGIN_BOTTOM_CM:.2f} см)")
        if abs(section.left_margin.cm - TARGET_MARGIN_LEFT_CM) > tolerance_cm:
            margin_errors.append(f"левое поле {section.left_margin.cm:.2f} см (нужно {TARGET_MARGIN_LEFT_CM:.2f} см)")
        if abs(section.right_margin.cm - TARGET_MARGIN_RIGHT_CM) > tolerance_cm:
            margin_errors.append(
                f"правое поле {section.right_margin.cm:.2f} см (нужно {TARGET_MARGIN_RIGHT_CM:.2f} см)")
        return margin_errors

    def _are_effective_fonts_same(self, run1_font_name, run1_font_size, run2_font_name, run2_font_size):
        size_matches = False
        if run1_font_size is None and run2_font_size is None:
            size_matches = True
        elif run1_font_size is not None and run2_font_size is not None and abs(run1_font_size - run2_font_size) < 0.1:
            size_matches = True
        return run1_font_name == run2_font_name and size_matches

    def analyze_docx(self, doc_path):
        self.results_text_edit.append("<br><i>Анализирую документ...</i>")
        QApplication.processEvents()

        structured_issues_report = []
        document_level_errors = []

        try:
            document = Document(doc_path)
            self.results_text_edit.append(f"Открыт документ. Абзацев: {len(document.paragraphs)}.")

            margin_errors = self._check_document_margins(document)
            if margin_errors:
                for err in margin_errors:
                    document_level_errors.append(f"Поля документа: {err}")
                    self.results_text_edit.append(
                        f"<font color='purple'><b>ПОЛЯ ДОКУМЕНТА:</b> {self._escape_html(err)}</font>")
            QApplication.processEvents()

            for p_idx, para in enumerate(document.paragraphs):
                if not para.text.strip():
                    continue

                para_preview = para.text[:35].strip().replace('\n', ' ')
                if len(para.text) > 35: para_preview += "..."

                current_para_report_data = {
                    'paragraph_index': p_idx + 1,
                    'paragraph_preview': para_preview,
                    'general_errors': [],
                    'logical_runs': []
                }
                para_has_any_issue = False
                para_display_header_shown_in_main_log = False

                fragment_error_counter_for_this_paragraph = 1

                para_formatting_errors = self._check_paragraph_formatting(para)
                if para_formatting_errors:
                    if not para_display_header_shown_in_main_log:
                        self.results_text_edit.append(
                            f"<br><b>--- Абзац №{p_idx + 1} (текст: \"{self._escape_html(para_preview)}\") ---</b>")
                        para_display_header_shown_in_main_log = True
                    para_has_any_issue = True
                    current_para_report_data['general_errors'].extend(para_formatting_errors)
                    for err in para_formatting_errors:
                        self.results_text_edit.append(
                            f"  <font color='blue'><b>ПАРАМЕТРЫ АБЗАЦА:</b> {self._escape_html(err)}</font>")

                current_para_style = para.style if (para.style and para.style.type == WD_STYLE_TYPE.PARAGRAPH) else None

                if para.runs:
                    logical_run_text_buffer = ""
                    if para.runs:
                        first_run_in_sequence = para.runs[0]
                        buffer_eff_font_name, buffer_eff_font_size = self._get_actual_font_for_run(
                            first_run_in_sequence, current_para_style)
                    else:
                        buffer_eff_font_name, buffer_eff_font_size = None, None

                    for run_idx, current_run in enumerate(para.runs):
                        current_run_eff_name, current_run_eff_size = self._get_actual_font_for_run(current_run,
                                                                                                   current_para_style)

                        is_last_run_in_paragraph = (run_idx == len(para.runs) - 1)
                        formatting_changed = not self._are_effective_fonts_same(buffer_eff_font_name,
                                                                                buffer_eff_font_size,
                                                                                current_run_eff_name,
                                                                                current_run_eff_size)

                        if formatting_changed:
                            if logical_run_text_buffer.strip():
                                font_errors = self._check_font_and_size(logical_run_text_buffer, buffer_eff_font_name,
                                                                        buffer_eff_font_size)
                                fragment_num = None
                                if font_errors:
                                    fragment_num = fragment_error_counter_for_this_paragraph
                                    fragment_error_counter_for_this_paragraph += 1
                                current_para_report_data['logical_runs'].append({
                                    'text': logical_run_text_buffer,
                                    'has_font_errors': bool(font_errors),
                                    'error_details': font_errors,
                                    'fragment_number': fragment_num
                                })
                                if font_errors: para_has_any_issue = True

                            logical_run_text_buffer = current_run.text
                            buffer_eff_font_name, buffer_eff_font_size = current_run_eff_name, current_run_eff_size
                        else:
                            logical_run_text_buffer += current_run.text

                        if is_last_run_in_paragraph:
                            if logical_run_text_buffer.strip():
                                font_errors = self._check_font_and_size(logical_run_text_buffer, buffer_eff_font_name,
                                                                        buffer_eff_font_size)
                                fragment_num = None
                                if font_errors:
                                    fragment_num = fragment_error_counter_for_this_paragraph
                                current_para_report_data['logical_runs'].append({
                                    'text': logical_run_text_buffer,
                                    'has_font_errors': bool(font_errors),
                                    'error_details': font_errors,
                                    'fragment_number': fragment_num
                                })
                                if font_errors: para_has_any_issue = True

                    for lr_data in current_para_report_data['logical_runs']:
                        if lr_data['has_font_errors']:
                            if not para_display_header_shown_in_main_log:
                                self.results_text_edit.append(
                                    f"<br><b>--- Абзац №{p_idx + 1} (текст: \"{self._escape_html(para_preview)}\") ---</b>")
                                para_display_header_shown_in_main_log = True

                            run_preview_short = lr_data['text'].strip()[:20]
                            if len(lr_data['text'].strip()) > 20: run_preview_short += "..."
                            frag_num_str = f" (фрагмент #{lr_data['fragment_number']})" if lr_data[
                                                                                               'fragment_number'] is not None else ""
                            self.results_text_edit.append(
                                f"  <font color='red'><b>ШРИФТ/РАЗМЕР (фрагмент \"{self._escape_html(run_preview_short)}\"{frag_num_str}):</b> {'; '.join(lr_data['error_details'])}</font>")

                if para_has_any_issue:
                    structured_issues_report.append(current_para_report_data)

                if para_display_header_shown_in_main_log:
                    QApplication.processEvents()

            # --- ФОРМИРОВАНИЕ ИТОГОВОГО ОТЧЕТА с использованием таблиц ---
            self.results_text_edit.append(
                "<hr><p style='margin-top:5px; margin-bottom:2px;'><b>--- ИТОГ ПРОВЕРКИ ---</b></p>")
            total_issues_found = len(document_level_errors) + len(structured_issues_report)

            if total_issues_found > 0:
                self.results_text_edit.append(
                    f"<p style='margin-top:2px; margin-bottom:5px;'><font color='red' size='+1'><b>Обнаружены несоответствия:</b></font></p>")
                if document_level_errors:
                    self.results_text_edit.append(
                        "<p style='margin-top:3px; margin-bottom:1px;'><font color='purple' size='+0'><b>Замечания по документу в целом:</b></font></p>")
                    for err_text in document_level_errors:
                        self.results_text_edit.append(
                            f"<p style='margin-left:10px; margin-top:1px; margin-bottom:1px;'>- <font size='+0'>{self._escape_html(err_text)}</font></p>")

                if structured_issues_report:
                    self.results_text_edit.append(
                        "<p style='margin-top:8px; margin-bottom:2px;'><font color='red' size='+0'><b>Проблемные абзацы:</b></font></p>")

                    for issue_data in structured_issues_report:
                        para_idx = issue_data['paragraph_index']
                        para_preview_text = issue_data['paragraph_preview']

                        table_html = "<table width='100%' cellspacing='0' cellpadding='2' style='margin-top: 5px; border-top: 1px solid #eee;'>"
                        table_html += f"<tr><td style='padding-left: 0px;'><h4 style='margin: 2px 0;'>Абзац №{para_idx} (начинается с: \"{self._escape_html(para_preview_text)}\")</h4></td></tr>"

                        if issue_data['general_errors']:
                            for err in issue_data['general_errors']:
                                table_html += f"<tr><td style='padding-left: 10px;'><font color='blue' size='+0'>- Параметры абзаца: {self._escape_html(err)}</font></td></tr>"

                        has_font_size_errors_in_para_summary = any(
                            lr_data['has_font_errors'] for lr_data in issue_data['logical_runs'])
                        if has_font_size_errors_in_para_summary:
                            table_html += f"<tr><td style='padding-left: 10px;'><font size='+0'><u>- Ошибки шрифта/размера в тексте:</u></font></td></tr>"

                            highlighted_paragraph_html = ""
                            for lr_data in issue_data['logical_runs']:
                                escaped_run_text = self._escape_html(lr_data['text'])
                                if lr_data['has_font_errors'] and lr_data['fragment_number'] is not None:
                                    # ИЗМЕНЕНО: Маркер перед текстом и другой цвет
                                    error_marker = f" <font color='#0000CC'><b>({lr_data['fragment_number']})</b></font> "
                                    highlighted_paragraph_html += f"{error_marker}<font color='red'><b>{escaped_run_text}</b></font>"
                                elif lr_data['has_font_errors']:
                                    highlighted_paragraph_html += f"<font color='red'><b>{escaped_run_text}</b></font>"
                                else:
                                    highlighted_paragraph_html += escaped_run_text
                            table_html += f"<tr><td style='padding-left: 25px;'><div style='border:1px solid #ddd; padding:4px; margin:0 0 2px 0; background-color:#fff;'>{highlighted_paragraph_html}</div></td></tr>"

                            details_exist = any(lr_data['has_font_errors'] and lr_data['error_details'] for lr_data in
                                                issue_data['logical_runs'])
                            if details_exist:
                                # ИЗМЕНЕНО: Увеличен размер шрифта для заголовка деталей
                                table_html += f"<tr><td style='padding-left: 25px;'><font size='+0'>Детали по фрагментам с ошибками:</font></td></tr>"
                                for lr_data in issue_data['logical_runs']:
                                    if lr_data['has_font_errors'] and lr_data['error_details'] and lr_data[
                                        'fragment_number'] is not None:
                                        # ИЗМЕНЕНО: Увеличен размер шрифта для самих деталей
                                        table_html += (
                                            f"<tr><td style='padding-left: 30px;'><font size='+0'>- Фрагмент ({lr_data['fragment_number']}): {self._escape_html('; '.join(lr_data['error_details']))}</font></td></tr>"
                                        )
                        table_html += "</table>"
                        self.results_text_edit.append(table_html)
            else:
                self.results_text_edit.append(
                    "<p style='margin-top:5px;'><font color='green' size='+1'><b>Отлично! Несоответствий по проверяемым параметрам не найдено.</b></font></p>")
            self.results_text_edit.append("<br><i>Анализ завершен.</i>")

        except Exception as e:
            error_msg = f"Произошла непредвиденная ошибка: {str(e)}"
            self.results_text_edit.append(f"<br><font color='red'><b>{error_msg}</b></font>")
            print(f"Критическая ошибка при анализе: {e}")
            traceback.print_exc()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    checker_window = DocFormatChecker()
    sys.exit(app.exec_())
